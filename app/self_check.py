"""Fault-tolerant command-line diagnostics used by users and packaged builds."""
from __future__ import print_function
import importlib
import copy
import json
import os
import platform
import pkgutil
import shutil
import subprocess
import sys
import tempfile
import time
import traceback

RESET = '\033[0m'
BOLD = '\033[1m'
GREEN = '\033[32m'
RED = '\033[31m'
CYAN = '\033[36m'


def _color(code, text):
    return code + text + RESET


def _check_python():
    if sys.version_info[:2] < (3, 7):
        raise RuntimeError('Python 3.7 or newer is required')
    return platform.python_version()


def _check_import(name):
    module = importlib.import_module(name)
    return getattr(module, '__version__', 'imported')


def _check_java():
    path = shutil.which('java')
    if not path:
        raise RuntimeError('java is not on PATH (PlantUML rendering will be unavailable)')
    return path


def _check_plantuml_jar():
    roots = [getattr(sys, '_MEIPASS', ''), os.path.dirname(os.path.dirname(os.path.abspath(__file__)))]
    for root in roots:
        path = os.path.join(root, 'docs', 'plantuml.jar') if root else ''
        if path and os.path.isfile(path):
            return path
    raise RuntimeError('docs/plantuml.jar is missing')


def _check_qt_application():
    os.environ.setdefault('QT_QPA_PLATFORM', 'offscreen')
    from PyQt5.QtWidgets import QApplication
    app = QApplication.instance() or QApplication([])
    app.processEvents()
    return 'event processing OK'


def _qt_platform_name():
    try:
        from PyQt5.QtGui import QGuiApplication
        return str(QGuiApplication.platformName() or '<not-created>')
    except BaseException:
        return '<unavailable>'


def _pyfcstm_module_names():
    import pyfcstm
    names = []
    for module in pkgutil.walk_packages(pyfcstm.__path__, pyfcstm.__name__ + '.'):
        # Command entry points parse argv when executed and are covered by
        # their underlying libraries instead of being imported as scripts.
        if not module.name.endswith('.__main__'):
            names.append(module.name)
    return sorted(names)


def _check_pyfcstm_roundtrip():
    from pyfcstm.dsl import parse_with_grammar_entry
    from pyfcstm.model import parse_dsl_node_to_state_machine
    source = 'state Smoke { state Idle; [*] -> Idle; Idle -> [*]; }'
    ast_node = parse_with_grammar_entry(source, entry_name='state_machine_dsl')
    machine = parse_dsl_node_to_state_machine(ast_node)
    ast_text = str(machine.to_ast_node())
    plantuml = machine.to_plantuml()
    if machine.root_state.name != 'Smoke' or 'state Smoke' not in ast_text or '@startuml' not in plantuml:
        raise RuntimeError('pyfcstm roundtrip returned incomplete output')
    return '{} AST chars, {} PlantUML chars'.format(len(ast_text), len(plantuml))


def _check_z3_integer_sat():
    import z3
    value = z3.Int('fcstm_self_check')
    sat_solver = z3.Solver()
    sat_solver.add(value > 5, value < 10)
    if sat_solver.check() != z3.sat or not 5 < sat_solver.model()[value].as_long() < 10:
        raise RuntimeError('Z3 integer SAT/model check failed')
    return 'integer SAT/model OK'


def _check_z3_unsat():
    import z3
    value = z3.Int('fcstm_unsat')
    unsat_solver = z3.Solver()
    unsat_solver.add(value > 2, value < 2)
    if unsat_solver.check() != z3.unsat:
        raise RuntimeError('Z3 UNSAT check failed')
    return 'UNSAT OK'


def _check_z3_real():
    import z3
    real = z3.Real('fcstm_real')
    real_solver = z3.Solver()
    real_solver.add(real * 3 == 1)
    if real_solver.check() != z3.sat or str(real_solver.model()[real]) != '1/3':
        raise RuntimeError('Z3 real arithmetic check failed')
    return 'exact real 1/3 OK'


def _check_z3_bitvector():
    import z3
    bits = z3.BitVec('fcstm_bits', 8)
    bit_solver = z3.Solver()
    bit_solver.add((bits & 0x0f) == 0x0a, bits == 0x2a)
    if bit_solver.check() != z3.sat or bit_solver.model()[bits].as_long() != 0x2a:
        raise RuntimeError('Z3 bit-vector check failed')
    return '8-bit model 0x2a OK'


def _check_z3_optimize():
    import z3
    optimum = z3.Int('fcstm_optimum')
    optimizer = z3.Optimize()
    optimizer.add(optimum >= 0, optimum <= 20)
    optimizer.maximize(optimum)
    if optimizer.check() != z3.sat or optimizer.model()[optimum].as_long() != 20:
        raise RuntimeError('Z3 Optimize/maximize check failed')
    return 'Optimize maximum 20 OK'


def _check_qt_native_widgets():
    from PyQt5.Qsci import QsciScintilla
    from PyQt5.QtWidgets import QApplication
    import qtawesome
    import qtmodern.styles
    app = QApplication.instance() or QApplication([])
    editor = QsciScintilla()
    editor.setText('state Smoke;')
    icon = qtawesome.icon('fa5s.check')
    qtmodern.styles.dark(app)
    editor.show()
    app.processEvents()
    if editor.text() != 'state Smoke;' or icon.isNull():
        raise RuntimeError('Qt native widget/font integration failed')
    editor.close()
    return 'QScintilla + qtawesome + qtmodern OK'


def _check_qt_cjk_font():
    import hashlib
    from PyQt5 import QtCore, QtGui
    from PyQt5.QtWidgets import QApplication
    from app.utils.application_font import (
        APPLICATION_FONT_POINT_SIZE,
        EXPECTED_FAMILY,
        FONT_SHA256,
        bundled_font_path,
        install_application_font,
    )
    app = QApplication.instance() or QApplication([])
    data = bundled_font_path().read_bytes()
    if hashlib.sha256(data).hexdigest() != FONT_SHA256:
        raise RuntimeError('bundled CJK font SHA-256 mismatch')
    family = install_application_font(app)
    if (
        family != EXPECTED_FAMILY
        or app.font().family() != EXPECTED_FAMILY
        or app.font().pointSize() != APPLICATION_FONT_POINT_SIZE
    ):
        raise RuntimeError('bundled CJK font did not become the application font')
    image = QtGui.QImage(320, 80, QtGui.QImage.Format_ARGB32)
    image.fill(QtCore.Qt.white)
    painter = QtGui.QPainter(image)
    painter.setFont(QtGui.QFont(EXPECTED_FAMILY, 18))
    painter.setPen(QtCore.Qt.black)
    painter.drawText(image.rect(), QtCore.Qt.AlignCenter, '状态机动态验证')
    painter.end()
    dark_pixels = 0
    for y in range(image.height()):
        for x in range(image.width()):
            if QtGui.qGray(image.pixel(x, y)) < 220:
                dark_pixels += 1
    if dark_pixels < 200:
        raise RuntimeError('bundled CJK font produced no visible Chinese text')
    return '{} rendered {} dark pixels'.format(family, dark_pixels)


def _check_office_roundtrips():
    from io import BytesIO
    from docx import Document
    from openpyxl import Workbook, load_workbook
    xlsx = BytesIO()
    workbook = Workbook()
    workbook.active['A1'] = 'fcstm'
    workbook.save(xlsx)
    xlsx.seek(0)
    if load_workbook(xlsx).active['A1'].value != 'fcstm':
        raise RuntimeError('openpyxl XLSX roundtrip failed')
    docx = BytesIO()
    document = Document()
    document.add_paragraph('fcstm')
    document.save(docx)
    docx.seek(0)
    if Document(docx).paragraphs[0].text != 'fcstm':
        raise RuntimeError('python-docx roundtrip failed')
    return '{} XLSX bytes, {} DOCX bytes'.format(len(xlsx.getvalue()), len(docx.getvalue()))


def _plantuml_path():
    roots = [getattr(sys, '_MEIPASS', ''), os.path.dirname(os.path.dirname(os.path.abspath(__file__)))]
    for root in roots:
        path = os.path.join(root, 'docs', 'plantuml.jar') if root else ''
        if path and os.path.isfile(path):
            return path
    raise RuntimeError('docs/plantuml.jar is missing')


def _check_plantuml_render():
    java = shutil.which('java')
    if not java:
        raise RuntimeError('java is not on PATH')
    with tempfile.TemporaryDirectory() as directory:
        source = os.path.join(directory, 'smoke.puml')
        with open(source, 'w') as file:
            file.write('@startuml\n[*] --> Ready\n@enduml\n')
        process = subprocess.run([java, '-jar', _plantuml_path(), '-tpng', source],
                                 stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
        output = os.path.join(directory, 'smoke.png')
        if process.returncode != 0 or not os.path.isfile(output) or os.path.getsize(output) < 100:
            raise RuntimeError('PlantUML PNG render failed: {!r}'.format(process.stderr[-500:]))
        return '{} PNG bytes'.format(os.path.getsize(output))


def _check_pyfcstm_simulation():
    from pyfcstm.dsl import parse_with_grammar_entry
    from pyfcstm.model import parse_dsl_node_to_state_machine
    from pyfcstm.simulate import SimulationRuntime
    source = 'state Smoke { state Idle; [*] -> Idle; Idle -> [*]; }'
    machine = parse_dsl_node_to_state_machine(parse_with_grammar_entry(source, entry_name='state_machine_dsl'))
    runtime = SimulationRuntime(machine)
    return type(runtime).__name__ + ' constructed'


def _check_main_window():
    from PyQt5.QtWidgets import QApplication
    from app.widget import AppMainWindow
    app = QApplication.instance() or QApplication([])
    window = AppMainWindow()
    window.show()
    app.processEvents()
    if not window.isVisible():
        raise RuntimeError('main window did not become visible')
    window.close()
    return 'constructed, shown, events processed, closed'


_ACCEPTANCE_DSL = 'state Smoke { state Idle; [*] -> Idle; Idle -> [*]; }'


def _invoke_pyfcstm(arguments, files=None):
    from click.testing import CliRunner
    from pyfcstm.entry.cli import cli
    runner = CliRunner()
    with runner.isolated_filesystem():
        with open('acceptance.fcstm', 'w') as file:
            file.write(_ACCEPTANCE_DSL)
        result = runner.invoke(cli, arguments)
        if result.exit_code != 0:
            raise RuntimeError('pyfcstm {} failed: {}'.format(' '.join(arguments), result.output[-1000:]))
        outputs = {}
        for path in files or ():
            if not os.path.isfile(path) or os.path.getsize(path) == 0:
                raise RuntimeError('pyfcstm did not create non-empty ' + path)
            with open(path, 'rb') as file:
                outputs[path] = file.read()
        return result.output, outputs


def _check_invalid_diagnostics():
    from click.testing import CliRunner
    from pyfcstm.entry.cli import cli
    runner = CliRunner()
    with runner.isolated_filesystem():
        with open('invalid.fcstm', 'w') as file:
            file.write('state Broken { state ; }')
        result = runner.invoke(cli, ['inspect', '-i', 'invalid.fcstm', '--format', 'json'])
        if result.exit_code == 0 or not any(token in result.output.lower() for token in ('line', 'column', 'syntax', 'error')):
            raise RuntimeError('invalid DSL did not produce positional diagnostics: ' + result.output[-500:])
        return 'rejected with diagnostic output'


def _check_inspect(format_name):
    output, _ = _invoke_pyfcstm(['inspect', '-i', 'acceptance.fcstm', '--format', format_name, '--color', 'never'])
    if 'Smoke' not in output:
        raise RuntimeError(format_name + ' inspect omitted machine name')
    return '{} chars'.format(len(output))


def _check_template(template_name):
    from click.testing import CliRunner
    from pyfcstm.entry.cli import cli
    runner = CliRunner()
    with runner.isolated_filesystem():
        with open('acceptance.fcstm', 'w') as file:
            file.write(_ACCEPTANCE_DSL)
        result = runner.invoke(cli, ['generate', '-i', 'acceptance.fcstm', '--template', template_name,
                                     '-o', 'generated', '--clear'])
        if result.exit_code != 0:
            raise RuntimeError(template_name + ' generation failed: ' + result.output[-1000:])
        generated = []
        for root, _, names in os.walk('generated'):
            generated.extend(os.path.join(root, name) for name in names if os.path.getsize(os.path.join(root, name)) > 0)
        if not generated:
            raise RuntimeError(template_name + ' template produced no non-empty files')
        return '{} files'.format(len(generated))


def _check_pyfcstm_plantuml_cli():
    _, outputs = _invoke_pyfcstm(['plantuml', '-i', 'acceptance.fcstm', '-o', 'machine.puml'], ['machine.puml'])
    if b'@startuml' not in outputs['machine.puml']:
        raise RuntimeError('plantuml CLI output is invalid')
    return '{} bytes'.format(len(outputs['machine.puml']))


def _check_visualize(format_name):
    from app.application.graph_render import (
        GraphRenderService,
        _reject_unexpected_stderr,
    )
    from pyfcstm.model import load_state_machine_from_text
    previous_dot = os.environ.get('GRAPHVIZ_DOT')
    os.environ['GRAPHVIZ_DOT'] = os.path.join(
        tempfile.gettempdir(), 'fcstm-gui-definitely-missing-dot'
    )
    try:
        model = load_state_machine_from_text(_ACCEPTANCE_DSL)
        with tempfile.TemporaryDirectory(prefix='fcstm-smetana-self-check-') as directory:
            target = os.path.join(directory, 'machine.' + format_name)
            result = GraphRenderService().render(
                model.to_plantuml(),
                target,
                format_name,
                plantuml_jar=_plantuml_path(),
            )
            if result.engine != 'smetana' or not result.output_sha256:
                raise RuntimeError(format_name + ' visualization omitted Smetana provenance')
            if result.exit_code != 0:
                raise RuntimeError(
                    format_name + ' renderer execution failed: ' + result.stderr
                )
            _reject_unexpected_stderr(result.stderr)
            if 'Smoke' not in result.semantic_labels:
                raise RuntimeError(format_name + ' visualization omitted model semantics')
            return '{} bytes, source {}, svg {}'.format(
                result.size,
                result.source_sha256[:12],
                result.semantic_svg_sha256[:12],
            )
    finally:
        if previous_dot is None:
            os.environ.pop('GRAPHVIZ_DOT', None)
        else:
            os.environ['GRAPHVIZ_DOT'] = previous_dot


def _check_pyfcstm_visualize_cli(format_name):
    output_name = 'machine.' + format_name
    _, outputs = _invoke_pyfcstm(
        [
            'visualize', '-i', 'acceptance.fcstm', '-o', output_name,
            '--type', format_name, '--renderer', 'local',
            '--plantuml-jar', _plantuml_path(), '--no-open',
        ],
        [output_name],
    )
    magic = {'png': b'\x89PNG', 'svg': b'<?xml', 'pdf': b'%PDF'}[format_name]
    if not outputs[output_name].startswith(magic):
        raise RuntimeError(format_name + ' production CLI output has incorrect magic')
    return '{} bytes'.format(len(outputs[output_name]))


def _check_simulate_cli():
    output, _ = _invoke_pyfcstm(['simulate', '-i', 'acceptance.fcstm', '--execute', 'current', '--no-color'])
    if 'Idle' not in output and 'Smoke' not in output:
        raise RuntimeError('batch simulation produced no state output: ' + output[-500:])
    return '{} chars'.format(len(output))


def _check_pygments_highlight():
    from pygments import highlight
    from pygments.formatters import HtmlFormatter
    from pyfcstm.highlight.pygments_lexer import FcstmLexer
    output = highlight(_ACCEPTANCE_DSL, FcstmLexer(), HtmlFormatter())
    if 'Smoke' not in output or '<span' not in output:
        raise RuntimeError('FCSTM Pygments highlighting failed')
    return '{} HTML chars'.format(len(output))


_BEHAVIOR_SOURCE = '''
def int x = 0;
def int y = 1;
state Root {
    state A { during { x = x + 1; } }
    state B;
    [*] -> A;
    A -> B :: Go effect { x = x + 5; }
    B -> [*] :: Stop;
}
'''


def _check_loader_text():
    from pyfcstm.model import load_state_machine_from_text
    model = load_state_machine_from_text(_BEHAVIOR_SOURCE)
    if model.root_state.name != 'Root':
        raise RuntimeError('text loader returned the wrong root state')
    return 'root Root loaded from text'


def _check_loader_file():
    from pyfcstm.model import load_state_machine_from_file
    with tempfile.TemporaryDirectory() as directory:
        path = os.path.join(directory, 'model.fcstm')
        with open(path, 'w', encoding='utf-8') as stream:
            stream.write(_BEHAVIOR_SOURCE)
        model = load_state_machine_from_file(path)
    if model.root_state.name != 'Root':
        raise RuntimeError('file loader returned the wrong root state')
    return 'root Root loaded from file'


def _check_loader_syntax_failure():
    from pyfcstm.model import load_state_machine_from_text
    try:
        load_state_machine_from_text('state Broken { state ; }')
    except BaseException as error:
        text = str(error).lower()
        if not any(item in text for item in ('line', 'column', 'syntax')):
            raise RuntimeError('syntax error omitted position: ' + str(error))
        return type(error).__name__ + ' with position'
    raise RuntimeError('invalid syntax was accepted')


def _check_loader_model_failure():
    from pyfcstm.model import load_state_machine_from_text
    try:
        load_state_machine_from_text(
            'state Root { state A; [*] -> Missing; }'
        )
    except BaseException as error:
        if 'Missing' not in str(error):
            raise RuntimeError('model validation omitted invalid target')
        return type(error).__name__ + ' for missing target'
    raise RuntimeError('invalid model structure was accepted')


def _check_inspect_warning():
    from pyfcstm.diagnostics.inspect import inspect_model
    from pyfcstm.model import load_state_machine_from_text
    model = load_state_machine_from_text(
        'state Root { state A; [*] -> A; }'
    )
    report = inspect_model(model)
    warnings = [item for item in report.diagnostics if item.severity == 'warning']
    if not warnings or not all(item.code and item.span for item in warnings):
        raise RuntimeError('inspect warning/code/span not available')
    return '{} warning(s), first {}'.format(len(warnings), warnings[0].code)


def _formula_check(kind, text, expected_valid):
    from app.application.formulas import (
        FormulaKind,
        FormulaValidationRequest,
        FormulaValidationService,
        FormulaValidationStatus,
    )
    result = FormulaValidationService().validate(
        FormulaValidationRequest(
            kind=FormulaKind(kind),
            text=text,
            source_revision=1,
            request_token='self-check-' + kind,
            variable_definitions='def int x = 0;\ndef int y = 1;',
        )
    )
    expected = (
        FormulaValidationStatus.VALID
        if expected_valid
        else FormulaValidationStatus.INVALID
    )
    if result.status is not expected:
        raise RuntimeError(
            '{} formula expected {}, got {}: {}'.format(
                kind, expected.value, result.status.value, result.message
            )
        )
    if not expected_valid and result.location is None:
        raise RuntimeError(kind + ' invalid result omitted location')
    return '{} at revision {}'.format(result.status.value, result.source_revision)


def _simulation_session():
    from app.application.simulation import SimulationService
    service = SimulationService()
    return service, service.start(
        _BEHAVIOR_SOURCE,
        source_uri='self-check://simulation',
        source_revision=1,
        dependency_fingerprint='self-check-model',
    )


def _check_simulation_initialization():
    service, session = _simulation_session()
    snapshot = session.snapshot()
    if snapshot.cycle != 0 or snapshot.vars != {'x': 0, 'y': 1}:
        raise RuntimeError('simulation initialization snapshot is wrong')
    return type(session.runtime).__name__ + ' cycle 0'


def _check_simulation_cycles():
    service, session = _simulation_session()
    first = service.cycle(session)
    second = service.cycle(session, events='Go')
    if (
        first.snapshot.state_path != ('Root', 'A')
        or second.snapshot.state_path != ('Root', 'B')
        or second.snapshot.vars['x'] != 6
        or second.consumed_events != ('Root.A.Go',)
    ):
        raise RuntimeError('multiple simulation cycles returned wrong state/variables')
    return '2 cycles, Root.B, x=6'


def _check_simulation_end():
    service, session = _simulation_session()
    service.cycle(session)
    service.cycle(session, events='Go')
    ended = service.cycle(session, events='Stop')
    if not ended.snapshot.ended or ended.snapshot.state_path:
        raise RuntimeError('simulation did not reach terminal state')
    return 'terminal state reached at cycle {}'.format(ended.snapshot.cycle)


def _check_simulation_exception_rollback():
    from app.application.simulation import SimulationService
    source = '''
def int x = 0;
def int y = 1;
state Root {
 state A;
 state B;
 [*] -> A;
 A -> B :: Boom effect { x = x + 1; y = 1 / 0; }
}
'''
    service = SimulationService()
    session = service.start(source, 'self-check://rollback', 1, 'rollback')
    service.cycle(session)
    before = session.snapshot()
    result = service.cycle(session, events='Boom')
    if (
        result.error is None
        or result.error.cause_type != 'ZeroDivisionError'
        or result.rollback_preserved is not True
        or result.snapshot != before
    ):
        raise RuntimeError('simulation exception/cause/rollback evidence is wrong')
    return '{} caused by {}, rollback preserved'.format(
        result.error.type, result.error.cause_type
    )


def _check_dynamic_case(case_id):
    from app.application.dynamic_validation import DynamicValidationService
    report = DynamicValidationService().run_packaged_case(case_id)
    if report.status != 'passed' or not report.steps:
        raise RuntimeError(case_id + ' did not pass')
    return '{} step(s), scenario {}'.format(
        len(report.steps), report.scenario_sha256[:12]
    )


def _mutated_dynamic_report():
    from app.application.dynamic_validation import DynamicValidationService
    service = DynamicValidationService()
    case_id = 'design_validation_failure_multilevel_transition'
    scenario_path = service.resource_dir / (case_id + '.json')
    payload = json.loads(scenario_path.read_text(encoding='utf-8'))
    payload = copy.deepcopy(payload)
    payload['case_id'] = 'self_check_mutation'
    payload['steps'][-1]['expected']['state'] = 'Root.Mutated'
    with tempfile.TemporaryDirectory() as directory:
        model_name = payload['model_file']
        source = service.resource_dir / model_name
        target = os.path.join(directory, model_name)
        with open(target, 'wb') as stream:
            stream.write(source.read_bytes())
        return service.run_scenario(payload, base_dir=directory)


def _check_dynamic_mutation():
    report = _mutated_dynamic_report()
    if report.status != 'mismatch' or not any(step.diffs for step in report.steps):
        raise RuntimeError('dynamic mutation was not detected')
    return 'expected-state mutation detected'


def _check_dynamic_restore():
    _mutated_dynamic_report()
    return _check_dynamic_case('design_validation_failure_multilevel_transition')


def _check_dynamic_provenance():
    from app.application.dynamic_validation import DynamicValidationService
    report = DynamicValidationService().verify_packaged_provenance()
    if report.status != 'passed' or len(report.resources) != 8:
        raise RuntimeError('dynamic validation provenance mismatch')
    return '8 resource hashes match upstream provenance'


def _check_template_list():
    from app.application.generation import GenerationService
    names = [item.name for item in GenerationService().list_templates()]
    expected = ['c', 'c_poll', 'cpp', 'cpp_poll', 'python']
    if names != expected:
        raise RuntimeError('template inventory mismatch: {!r}'.format(names))
    return ', '.join(names)


def _check_template_info(template_name):
    from pyfcstm.template import get_template_info
    info = get_template_info(template_name)
    if info.get('name') != template_name or not info.get('language') or not info.get('description'):
        raise RuntimeError(template_name + ' metadata is incomplete')
    return '{} / {}'.format(info['title'], info['language'])


def _check_template_extract(template_name):
    from pyfcstm.template import extract_template
    with tempfile.TemporaryDirectory() as directory:
        path = extract_template(template_name, directory)
        if not os.path.isfile(os.path.join(path, 'config.yaml')):
            raise RuntimeError(template_name + ' extraction omitted config.yaml')
        count = sum(len(names) for _, _, names in os.walk(path))
    if count < 2:
        raise RuntimeError(template_name + ' extracted too few files')
    return '{} extracted files'.format(count)


def _checks():
    checks = [('python runtime', _check_python)]
    for name in ('PyQt5', 'qtpy', 'qtawesome', 'qtmodern', 'openpyxl', 'docx', 'pyfcstm'):
        checks.append(('import ' + name, lambda name=name: _check_import(name)))
    checks.extend([('java executable', _check_java), ('plantuml.jar', _check_plantuml_jar),
                   ('loader text success', _check_loader_text),
                   ('loader file success', _check_loader_file),
                   ('loader syntax failure position', _check_loader_syntax_failure),
                   ('loader model assembly failure', _check_loader_model_failure),
                   ('inspect warning/code/span', _check_inspect_warning),
                   ('pyfcstm DSL/model/PlantUML roundtrip', _check_pyfcstm_roundtrip),
                   ('simulation runtime construction', _check_pyfcstm_simulation),
                   ('simulation initialization', _check_simulation_initialization),
                   ('simulation multiple cycles state variables', _check_simulation_cycles),
                   ('simulation terminal state', _check_simulation_end),
                   ('simulation exception cause rollback', _check_simulation_exception_rollback),
                   ('pyfcstm invalid syntax diagnostics', _check_invalid_diagnostics),
                   ('pyfcstm inspect human', lambda: _check_inspect('human')),
                   ('pyfcstm inspect json', lambda: _check_inspect('json')),
                   ('pyfcstm PlantUML CLI', _check_pyfcstm_plantuml_cli),
                   ('Smetana SVG without Graphviz', lambda: _check_visualize('svg')),
                   ('Smetana PNG without Graphviz', lambda: _check_visualize('png')),
                   ('Smetana PDF without Graphviz', lambda: _check_visualize('pdf')),
                   ('pyfcstm visualize SVG CLI', lambda: _check_pyfcstm_visualize_cli('svg')),
                   ('pyfcstm visualize PNG CLI', lambda: _check_pyfcstm_visualize_cli('png')),
                   ('pyfcstm visualize PDF CLI', lambda: _check_pyfcstm_visualize_cli('pdf')),
                   ('pyfcstm batch simulation CLI', _check_simulate_cli),
                   ('pyfcstm Pygments highlighting', _check_pygments_highlight),
                   ('Z3 integer SAT and model', _check_z3_integer_sat),
                   ('Z3 UNSAT', _check_z3_unsat),
                   ('Z3 exact real', _check_z3_real),
                   ('Z3 bit-vector', _check_z3_bitvector),
                   ('Z3 Optimize maximize', _check_z3_optimize),
                   ('Qt application', _check_qt_application),
                   ('Qt native widgets and assets', _check_qt_native_widgets),
                   ('Qt bundled CJK font rendering', _check_qt_cjk_font),
                   ('XLSX and DOCX roundtrips', _check_office_roundtrips),
                   ('PlantUML Java PNG render', _check_plantuml_render),
                   ('main GUI window lifecycle', _check_main_window),
                   ('packaged template inventory', _check_template_list),
                   ('dynamic provenance resource hashes', _check_dynamic_provenance),
                   ('dynamic mutation mismatch', _check_dynamic_mutation),
                   ('dynamic restored resource rerun', _check_dynamic_restore)])
    for kind, valid_text, invalid_text in (
        ('logical', 'x > 0 && y < 3', 'x +'),
        ('numeric', 'x * 2 + 1', 'x > 0'),
        ('effect', 'x = x + 1;', 'x = ;'),
        ('lifecycle', 'x = x + 1;', 'x = x + 1;\nx = ;'),
    ):
        checks.append((
            'formula {} valid'.format(kind),
            lambda kind=kind, text=valid_text: _formula_check(kind, text, True),
        ))
        checks.append((
            'formula {} invalid'.format(kind),
            lambda kind=kind, text=invalid_text: _formula_check(kind, text, False),
        ))
    for case_id in (
        'design_evented_pseudo_chain_invalid_then_valid',
        'design_validation_failure_multilevel_transition',
        'expression_failure_transition_guard_raises_expression_error',
        'pseudo_self_loop_step_limit_raises_dfs_error',
    ):
        checks.append((
            'dynamic case ' + case_id,
            lambda case_id=case_id: _check_dynamic_case(case_id),
        ))
    for template_name in ('python', 'c', 'c_poll', 'cpp', 'cpp_poll'):
        checks.append(('pyfcstm template info ' + template_name,
                       lambda template_name=template_name: _check_template_info(template_name)))
        checks.append(('pyfcstm extract template ' + template_name,
                       lambda template_name=template_name: _check_template_extract(template_name)))
        checks.append(('pyfcstm generate template ' + template_name,
                       lambda template_name=template_name: _check_template(template_name)))
    try:
        module_names = _pyfcstm_module_names()
    except BaseException as exc:
        checks.append(('discover pyfcstm modules', lambda exc=exc: (_ for _ in ()).throw(exc)))
    else:
        for name in module_names:
            checks.append(('import ' + name, lambda name=name: _check_import(name)))
    return checks


def _result_kind(name):
    return 'module-closure' if name.startswith('import ') else 'behavior'


def _result_group(name):
    lowered = name.lower()
    for group in (
        'dynamic', 'simulation', 'formula', 'z3', 'template', 'loader',
        'inspect', 'plantuml', 'qt', 'office', 'gui',
    ):
        if group in lowered:
            return group
    return 'runtime'


def _write_json_report(path, payload):
    target = os.path.abspath(path)
    parent = os.path.dirname(target)
    if parent:
        os.makedirs(parent, exist_ok=True)
    temporary = target + '.tmp'
    try:
        with open(temporary, 'w', encoding='utf-8') as stream:
            json.dump(payload, stream, ensure_ascii=False, sort_keys=True, indent=2)
            stream.write('\n')
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary, target)
    finally:
        try:
            os.remove(temporary)
        except OSError:
            pass


def run_self_check(json_report=None):
    try:
        if os.name == 'nt':
            import colorama
            colorama.just_fix_windows_console()
    except BaseException:
        pass
    started_at = time.time()
    results = []
    try:
        checks = _checks()
    except BaseException as exc:
        checks = []
        results.append({
            'name': 'check discovery',
            'group': 'runtime',
            'kind': 'behavior',
            'status': 'failed',
            'duration_ms': 0,
            'detail': repr(exc),
        })
    print(_color(CYAN, 'fcstm-gui self-check: running {} checks'.format(len(checks))))
    for index, (name, check) in enumerate(checks, 1):
        check_started = time.time()
        try:
            detail = check()
            results.append({
                'name': name,
                'group': _result_group(name),
                'kind': _result_kind(name),
                'status': 'passed',
                'duration_ms': int((time.time() - check_started) * 1000),
                'detail': str(detail),
            })
            print('[{:02d}/{:02d}] {}: {} ({})'.format(index, len(checks), name, _color(GREEN, 'OK'), detail))
        except BaseException as exc:
            results.append({
                'name': name,
                'group': _result_group(name),
                'kind': _result_kind(name),
                'status': 'failed',
                'duration_ms': int((time.time() - check_started) * 1000),
                'detail': repr(exc),
            })
            print('[{:02d}/{:02d}] {}: {} ({!r})'.format(index, len(checks), name, _color(RED, 'FAIL'), exc))
            traceback.print_exc()
    failures = [item for item in results if item['status'] == 'failed']
    passed = len(results) - len(failures)
    status = _color(GREEN if not failures else RED, 'PASSED' if not failures else 'FAILED')
    print(_color(CYAN, 'fcstm-gui self-check:') + ' {} OK / {} FAIL - {}'.format(passed, len(failures), BOLD + status + RESET))
    report = {
        'schema': 'fcstm-gui.self-check-report',
        'version': 1,
        'status': 'passed' if not failures else 'failed',
        'started_at': started_at,
        'duration_ms': int((time.time() - started_at) * 1000),
        'platform': {
            'system': platform.system(),
            'release': platform.release(),
            'machine': platform.machine(),
            'python': platform.python_version(),
            'frozen': bool(getattr(sys, 'frozen', False)),
            'qt_platform': _qt_platform_name(),
        },
        'counts': {
            'total': len(results),
            'passed': passed,
            'failed': len(failures),
            'module_closure': sum(item['kind'] == 'module-closure' for item in results),
            'behavior': sum(item['kind'] == 'behavior' for item in results),
        },
        'results': results,
    }
    if json_report:
        _write_json_report(json_report, report)
    return 1 if failures else 0


if __name__ == '__main__':
    sys.exit(run_self_check())
