from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from ..model import StateManager
from .ui_to_dsl import format_transition_item, state_manager_to_dsl
from pyfcstm.dsl import parse_with_grammar_entry
from pyfcstm.model import parse_dsl_node_to_state_machine


def _set_cell(cell, value, *, bold=False, centered=False):
    cell.text = "" if value is None else str(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        if centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "宋体"
            run.font.size = Pt(9)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _shade_cell(cell, fill="D9EAF7"):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_header(row, values):
    for cell, value in zip(row.cells, values):
        _set_cell(cell, value, bold=True, centered=True)
        _shade_cell(cell)


def _lifecycle_text(state, lifecycle_type):
    lines = []
    for item in state.lifecycle:
        if item.get("type") != lifecycle_type:
            continue
        if item.get("is_abstract"):
            text = "abstract {}".format(item.get("name", "")).strip()
        else:
            parts = []
            if item.get("name"):
                parts.append(str(item["name"]))
            if item.get("action"):
                parts.append(str(item["action"]))
            text = "\n".join(parts)
        if item.get("comment"):
            text = "{}\n// {}".format(text, item["comment"]).strip()
        lines.append(text)
    return "\n\n".join(lines)


def _transition_type(transition):
    source = str(transition.get("source", "")).strip()
    target = str(transition.get("target", "")).strip()
    if source.startswith("!"):
        return "强制迁移"
    if source == "[*]":
        return "初始迁移"
    if target == "[*]":
        return "终止迁移"
    return "普通迁移"


def export_statechart_to_word(state_manager: StateManager, file_path: str):
    """将状态机整体信息导出为状态机、状态和迁移三张 Word 表格。"""
    if state_manager is None or state_manager.root_state is None:
        raise ValueError("状态机为空，无法导出 Word 文档")

    # 保留完整 DSL 校验边界，避免把不一致的界面投影导出为正式文档。
    dsl_content = state_manager_to_dsl(state_manager)
    ast_node = parse_with_grammar_entry(
        dsl_content, entry_name="state_machine_dsl"
    )
    parse_dsl_node_to_state_machine(ast_node)

    states = state_manager.get_all_states()
    transitions = [
        (state, transition)
        for state in states
        for transition in state.transitions
    ]
    forced_count = sum(
        str(transition.get("source", "")).strip().startswith("!")
        for _, transition in transitions
    )
    variable_lines = [
        line.strip()
        for line in state_manager.variable_definitions.splitlines()
        if line.strip()
    ]
    maximum_level = max(
        (len(state.get_full_path().split(".")) for state in states),
        default=0,
    )

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height,
        section.page_width,
    )
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = document.add_heading("状态机导出报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("一、状态机表格", level=1)
    summary_rows = (
        ("状态机名称", state_manager.root_state.name),
        ("根状态路径", state_manager.root_state.get_full_path()),
        ("导出文件", Path(file_path).name),
        ("变量定义", state_manager.variable_definitions),
        ("变量定义数", len(variable_lines)),
        ("状态总数", len(states)),
        ("迁移总数", len(transitions)),
        ("普通/初始/终止迁移数", len(transitions) - forced_count),
        ("强制迁移数", forced_count),
        ("最大状态层级", maximum_level),
    )
    machine_table = document.add_table(rows=len(summary_rows), cols=2)
    machine_table.style = "Table Grid"
    for row, (label, value) in zip(machine_table.rows, summary_rows):
        _set_cell(row.cells[0], label, bold=True, centered=True)
        _shade_cell(row.cells[0])
        _set_cell(row.cells[1], value)

    document.add_paragraph()
    document.add_heading("二、状态表格", level=1)
    state_headers = (
        "序号",
        "状态路径",
        "状态名称",
        "父状态",
        "状态类型",
        "子状态数",
        "进入动作",
        "执行中动作",
        "退出动作",
    )
    state_table = document.add_table(rows=1, cols=len(state_headers))
    state_table.style = "Table Grid"
    _set_header(state_table.rows[0], state_headers)
    for index, state in enumerate(states, 1):
        row = state_table.add_row()
        values = (
            index,
            state.get_full_path(),
            state.name,
            state.parent.get_full_path() if state.parent is not None else "无",
            "复合状态" if state.children else "简单状态",
            len(state.children),
            _lifecycle_text(state, "enter"),
            _lifecycle_text(state, "during"),
            _lifecycle_text(state, "exit"),
        )
        for column, value in enumerate(values):
            _set_cell(row.cells[column], value, centered=column in {0, 2, 4, 5})

    document.add_paragraph()
    document.add_heading("三、迁移表格", level=1)
    transition_headers = (
        "序号",
        "所属状态",
        "迁移类型",
        "源状态",
        "目标状态",
        "事件",
        "条件",
        "动作",
        "完整定义",
    )
    transition_table = document.add_table(
        rows=1, cols=len(transition_headers)
    )
    transition_table.style = "Table Grid"
    _set_header(transition_table.rows[0], transition_headers)
    for index, (state, transition) in enumerate(transitions, 1):
        definition = format_transition_item(transition).strip()
        if definition and not definition.endswith(";"):
            definition += ";"
        row = transition_table.add_row()
        values = (
            index,
            state.get_full_path(),
            _transition_type(transition),
            transition.get("source", ""),
            transition.get("target", ""),
            transition.get("event", ""),
            transition.get("condition", ""),
            transition.get("action", ""),
            definition,
        )
        for column, value in enumerate(values):
            _set_cell(row.cells[column], value, centered=column in {0, 2})

    document.save(file_path)
