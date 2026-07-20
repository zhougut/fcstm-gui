from docx import Document

from app.model import State, StateManager
from app.utils.export_to_word import export_statechart_to_word


def test_export_to_word_uses_structured_state_data(tmp_path):
    root = State(
        "TrafficLight",
        transitions=[
            {"source": "[*]", "target": "Idle", "event": "", "condition": "", "action": ""},
            {"source": "! *", "target": "Idle", "event": "", "condition": "a > 10", "action": ""},
        ],
    )
    idle = State(
        "Idle",
        lifecycle=[{
            "type": "enter",
            "name": "",
            "action": "a = 1",
            "is_abstract": False,
            "comment": "",
        }],
    )
    manager = StateManager(root)
    manager.add_state(root, idle)
    manager.variable_definitions = "def int a = 0;"
    output_path = tmp_path / "traffic-light.docx"

    export_statechart_to_word(manager, str(output_path))

    document = Document(str(output_path))
    assert len(document.tables) == 3

    summary, states, transitions = document.tables
    summary_values = {
        row.cells[0].text: row.cells[1].text for row in summary.rows
    }
    assert summary_values["状态机名称"] == "TrafficLight"
    assert summary_values["状态总数"] == "2"
    assert summary_values["迁移总数"] == "2"
    assert summary_values["强制迁移数"] == "1"

    assert [cell.text for cell in states.rows[0].cells] == [
        "序号",
        "状态路径",
        "状态名称",
        "父状态",
        "状态类型",
        "子状态数",
        "进入动作",
        "执行中动作",
        "退出动作",
    ]
    state_rows = {row.cells[2].text: row for row in states.rows[1:]}
    assert state_rows["TrafficLight"].cells[3].text == "无"
    assert state_rows["Idle"].cells[3].text == "TrafficLight"
    assert state_rows["Idle"].cells[6].text == "a = 1"

    assert [cell.text for cell in transitions.rows[0].cells] == [
        "序号",
        "所属状态",
        "迁移类型",
        "源状态",
        "目标状态",
        "事件",
        "条件",
        "动作",
        "完整定义",
    ]
    transition_rows = transitions.rows[1:]
    assert transition_rows[0].cells[2].text == "初始迁移"
    assert transition_rows[0].cells[8].text == "[*] -> Idle;"
    assert transition_rows[1].cells[2].text == "强制迁移"
    assert transition_rows[1].cells[8].text == "! * -> Idle : if [a > 10];"
